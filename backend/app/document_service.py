import fitz
import os
from docx import Document
def extract_text_from_docx(docx_path: str) -> str:
    document = Document(docx_path)

    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    return "\n".join(paragraphs)
def extract_images_from_docx(docx_path: str, output_dir: str):
    document = Document(docx_path)

    os.makedirs(output_dir, exist_ok=True)

    image_paths = []
    image_count = 1

    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob

            image_path = os.path.join(
                output_dir,
                f"docx_image_{image_count}.png"
            )

            with open(image_path, "wb") as image_file:
                image_file.write(image_data)

            image_paths.append(image_path)
            image_count += 1

    return image_paths

    return "\n".join(paragraphs)
def convert_pdf_to_images(pdf_path: str, output_dir: str):
    pdf_document = fitz.open(pdf_path)

    image_paths = []

    os.makedirs(output_dir, exist_ok=True)

    for page_number in range(len(pdf_document)):
        page = pdf_document[page_number]

        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))

        image_path = os.path.join(
            output_dir,
            f"page_{page_number + 1}.png"
        )

        pix.save(image_path)

        image_paths.append(image_path)

    pdf_document.close()

    return image_paths