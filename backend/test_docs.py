from docx import Document

doc = Document("test_data/sample_text_prescription.docx")

print("Paragraph Count:", len(doc.paragraphs))

for i, p in enumerate(doc.paragraphs):
    print(f"Paragraph {i+1}:")
    print(repr(p.text))
    print("-" * 50)